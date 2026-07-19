# -*- coding: utf-8 -*-
import docx
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

RED   = RGBColor(0xC0,0x39,0x2B)
GREEN = RGBColor(0x1E,0x7A,0x46)
INK   = RGBColor(0x1F,0x29,0x33)
GRAY  = RGBColor(0x6B,0x76,0x83)
YAHEI = "Microsoft YaHei"
CAL   = "Calibri"

d = docx.Document()
sec = d.sections[0]
sec.left_margin  = Cm(3.17); sec.right_margin  = Cm(3.17)
sec.top_margin   = Cm(2.54); sec.bottom_margin = Cm(2.54)

def run(p, text, bold=False, size=11, color=INK, font=YAHEI):
    r = p.add_run(text)
    r.bold = bold
    r.font.size = Pt(size)
    r.font.color.rgb = color
    r.font.name = font
    # đảm bảo chữ Trung render đúng font (eastAsia)
    rPr = r._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = rPr.makeelement(qn('w:rFonts'), {}); rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), font)
    return r

def para(space_after=6, align=None):
    p = d.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)
    if align is not None: p.alignment = align
    return p

# ---------- Title ----------
p = para(2, WD_ALIGN_PARAGRAPH.CENTER)
run(p, "Nhận xét bài tập — Buổi 2: Lượng từ + ", True, 19, RED, YAHEI)
run(p, "一点儿/有点儿", True, 19, RED, YAHEI)
p = para(10, WD_ALIGN_PARAGRAPH.CENTER)
run(p, "Màu sắc  ·  Học viên: ______________   Ngày: __________", False, 11, GRAY, CAL)

# ---------- Tổng quan ----------
p = para(10)
run(p, "Tổng quan: ", True, 12, RED, CAL)
run(p, "Xuất sắc — ", False, 12, RED, CAL)
run(p, "22/22", True, 12, RED, CAL)
run(p, " câu khách quan đúng hết! Nắm rất chắc lượng từ (", False, 12, RED, CAL)
run(p, "本/支/口/件/个/双/张/杯", True, 12, RED, YAHEI)
run(p, ") và ", False, 12, RED, CAL)
run(p, "有点儿/一点儿", True, 12, RED, YAHEI)
run(p, ". Chỉ còn 1 ô nói bỏ trống (§7 câu 2) và nên tập thêm dấu câu.", False, 12, RED, CAL)

# ---------- Bảng điểm ----------
p = para(6); run(p, "Bảng điểm theo phần", True, 15, RED, YAHEI)

rows = [
    ("1. Nối chữ với nghĩa", "6/6 ✔", GREEN, "Đúng hết"),
    ("2. Điền lượng từ",     "4/4 ✔", GREEN, "本・支・口・件 chuẩn"),
    ("3. Đọc hiểu",          "3/3 ✔", GREEN, "1A 2A 3A — đúng hết"),
    ("4. Sắp xếp câu",       "3/3 ✔", GREEN, "Trật tự đúng (thiếu dấu 。)"),
    ("5. Dịch đặt câu",      "3/3 ✔", GREEN, "Dùng đúng từ cho sẵn"),
    ("6. Nghe chọn đáp án",  "3/3 ✔", GREEN, "1A 2A 3A — đúng hết"),
    ("7. Nghe & nhắc lại",   "2/3",   RED,   "Câu 1, 3 đúng; câu 2 chưa ghi"),
    ("8. Trả lời câu hỏi",   "2/2 ✔", GREEN, "Cả 2 câu tốt, có mở rộng"),
]
t = d.add_table(rows=1, cols=3)
t.style = "Table Grid"
hdr = ["Phần", "Kết quả", "Nhận xét"]
for ci, txt in enumerate(hdr):
    c = t.rows[0].cells[ci]
    c.paragraphs[0].text = ""
    run(c.paragraphs[0], txt, True, 11, RED, CAL)
for left, kq, kqcol, nx in rows:
    cells = t.add_row().cells
    run(cells[0].paragraphs[0], left, False, 11, INK, YAHEI)
    run(cells[1].paragraphs[0], kq,   True,  11, kqcol, CAL)
    run(cells[2].paragraphs[0], nx,   False, 11, RED,  YAHEI)
for col, w in zip(t.columns, [Cm(5.08)]*3):
    for c in col.cells: c.width = w
para(8)

# ---------- Chỗ cần chỉnh ----------
p = para(6); run(p, "Chỗ cần chỉnh (nhỏ — không trừ điểm)", True, 15, RED, YAHEI)

p = para(2)
run(p, "• Dấu câu   ", True, 12, INK, YAHEI)
run(p, "Em viết: ", False, 11, GRAY, CAL)
run(p, "我有三个苹果 / 这个杯子有点儿小 …", False, 13, INK, YAHEI)
p = para(8)
run(p, "→ Sửa: ", True, 11, RED, CAL)
run(p, "Thêm dấu 。 cuối mỗi câu tự viết (mục 4, 5, 7). Phần 书写 của HSK có tính dấu câu — vd ", False, 13, RED, YAHEI)
run(p, "我有三个苹果。", True, 13, RED, YAHEI)

# ---------- Còn bỏ trống ----------
p = para(6); run(p, "Còn bỏ trống — làm nốt cho đủ", True, 15, RED, YAHEI)
p = para(4)
run(p, "§7 câu 2  ", True, 12, INK, YAHEI)
run(p, "nghe rồi nhắc lại: ", False, 12, INK, CAL)
run(p, "请给我一双筷子。", True, 13, RED, YAHEI)
run(p, "  (luyện lượng từ 双)", False, 11, GRAY, CAL)
# (§8 câu 2 học viên đã trả lời — xem mục §8 bên dưới)

# ---------- §8 trả lời ----------
p = para(6); run(p, "§8 — Trả lời vậy được chưa? Được, cả 2 câu!", True, 15, RED, YAHEI)
p = para(2); run(p, "• 你喜欢什么颜色？", True, 12, INK, YAHEI)
p = para(2)
run(p, "Em viết: ", False, 11, GRAY, CAL)
run(p, "我喜欢粉红色。我的衣服和书包主要是粉红色的", False, 13, INK, YAHEI)
p = para(2)
run(p, "Nhận xét: ", True, 11, RED, CAL)
run(p, "Rất tốt — 2 câu, có chi tiết cá nhân (衣服和书包). Chỉ thiếu dấu 。 cuối.", False, 12, RED, YAHEI)
p = para(8)
run(p, "Hay hơn: ", True, 11, RED, CAL)
run(p, "我最喜欢粉红色，我的衣服和书包大部分都是粉红色的。", False, 14, RED, YAHEI)

p = para(2); run(p, "• 你家有几口人？", True, 12, INK, YAHEI)
p = para(2)
run(p, "Em viết: ", False, 11, GRAY, CAL)
run(p, "我家有四口人。包括爸爸、妈妈、我和我弟弟。我们家小小的，很漂亮。", False, 13, INK, YAHEI)
p = para(2)
run(p, "Nhận xét: ", True, 11, RED, CAL)
run(p, "Rất hay — đúng lượng từ 口, lại có mở rộng tả nhà (家小小的、很漂亮), rất tự nhiên. 包括 vượt HSK1 mà dùng đúng 👍. Chỉnh nhỏ: dấu câu kiểu Trung (。、) và 我和我弟弟 → 我和弟弟.", False, 12, RED, YAHEI)
p = para(8)
run(p, "Hay hơn: ", True, 11, RED, CAL)
run(p, "我家有四口人：爸爸、妈妈、我和弟弟。我们家小小的，很漂亮。", False, 14, RED, YAHEI)

# ---------- Điểm sáng ----------
p = para(6); run(p, "Điểm sáng — giữ và phát huy", True, 15, RED, YAHEI)
for txt_runs in [
    [("• Lượng từ dùng cực chắc: ", INK, CAL, False, 12),
     ("本/支/口/件/个/双/张/杯", INK, YAHEI, True, 12),
     (" — chọn đúng hết ở cả điền, sắp xếp lẫn dịch.", INK, CAL, False, 12)],
    [("• Phân biệt tốt ", INK, CAL, False, 12),
     ("有点儿", INK, YAHEI, True, 12),
     (" (than phiền) — ", INK, CAL, False, 12),
     ("这个杯子有点儿小", INK, YAHEI, True, 12),
     (" đúng ngữ cảnh.", INK, CAL, False, 12)],
    [("• Nghe & đọc hiểu đúng 100%. Câu trả lời màu sắc có ý riêng, rất tự nhiên.", INK, CAL, False, 12)],
]:
    p = para(2)
    for text, col, font, bold, sz in txt_runs:
        run(p, text, bold, sz, col, font)

# ---------- Nhắc chung ----------
p = para(0)
run(p, "Nhắc chung: ", True, 12, RED, CAL)
run(p, "(1) làm nốt §7.2 (请给我一双筷子。); (2) thêm dấu câu 。/、 kiểu Trung cuối câu tự viết; (3) giữ thói quen thêm chi tiết cá nhân khi trả lời (như câu màu sắc & gia đình).", False, 12, INK, CAL)

out = "output/hsk1/buoi2_luongtu_mausac/baitap/nhanxet.docx"
d.save(out)
print("SAVED", out)
