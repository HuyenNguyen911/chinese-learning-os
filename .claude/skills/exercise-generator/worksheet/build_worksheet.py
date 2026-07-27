#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
build_worksheet.py — Data-driven renderer: baitap JSON -> worksheet.docx + dapan.docx.

Dùng bởi skill exercise-generator (Giai đoạn B). Nhận 1 file JSON mô tả bài tập
(danh sách block: noi / dien_cho_trong / doc_hieu / sap_xep / dich_dat_cau /
nghe / noi_hskk / grammar_note / writing_prompt / dien_bieu_mau) và render ra:
  - hocsinh/worksheet.docx : bản cho học viên (KHÔNG có đáp án, KHÔNG có 听力文本)
  - dapan/dapan.docx       : đáp án + 听力文本 + gợi ý chấm
File nghe (audio/) do bước edge-tts sinh vào hocsinh/audio/ (cạnh worksheet).

Chạy:
    python build_worksheet.py <baitap.json> <out_dir>

Cần: python-docx. Không cần internet lúc render.
"""

import json
import sys
import shutil
import subprocess
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.opc.constants import RELATIONSHIP_TYPE as RT

DEFAULT_THEME = {
    "accent": "C0392B",
    "ink": "1F2933",
    "muted": "6B7683",
    "band": "F4F5F7",
    "accent_soft": "F7E4E1",
    "bg": "FFFFFF",
    "cjk_font": "Microsoft YaHei",
    "text_font": "Calibri",
}

BLANK = "＿＿＿＿"  # ô trống để học viên gõ (full-width để dễ thấy)


class WorksheetBuilder:
    def __init__(self, spec):
        self.spec = spec
        self.meta = spec.get("meta", {})
        self.theme = {**DEFAULT_THEME, **spec.get("theme", {})}

    # -- helpers -----------------------------------------------------------
    def _rgb(self, key):
        try:
            return RGBColor.from_string(self.theme[key])
        except KeyError:
            raise KeyError("unknown theme color %r; valid: %s" %
                           (key, sorted(self.theme)))

    def _run(self, p, text, size=12, color="ink", bold=False, italic=False,
             cjk=False):
        run = p.add_run(text)
        f = run.font
        f.size = Pt(size)
        f.bold = bold
        f.italic = italic
        f.color.rgb = self._rgb(color)
        f.name = self.theme["cjk_font"] if cjk else self.theme["text_font"]
        if cjk:
            run._element.rPr.rFonts.set(qn("w:eastAsia"), self.theme["cjk_font"])
        return run

    def _title_block(self, doc, text):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self._run(p, text, 20, color="accent", bold=True, cjk=True)

    def _block_header(self, doc, idx, title, instructions=None):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(14)
        self._run(p, "%d. " % idx, 15, color="accent", bold=True)
        self._run(p, title, 15, color="ink", bold=True, cjk=True)
        if instructions:
            ip = doc.add_paragraph()
            self._run(ip, instructions, 11, color="muted", italic=True, cjk=True)

    def _fill_cell(self, cell, text, size=12, color="ink", bold=False,
                   cjk=False, italic=False):
        p = cell.paragraphs[0]
        self._run(p, text, size, color=color, bold=bold, italic=italic, cjk=cjk)

    @staticmethod
    def _rotate(seq, by):
        # Deterministic reorder so worksheet order != source order, no RNG.
        by = by % len(seq) if seq else 0
        return seq[by:] + seq[:by]

    def _audio_link(self, paragraph, url, text, size=11, color="accent"):
        """Chèn hyperlink audio bấm được (Ctrl+Click mở MP3 bằng player mặc
        định). `url` là đường dẫn tương đối (vd audio/nghe-1.mp3) — thư mục
        audio/ phải nằm cạnh file .docx thì link mới trỏ đúng.

        File cục bộ (không phải http) dùng dấu `\\` để Word Windows hiểu là
        đường dẫn file, không nhầm sang URL web."""
        target = url if "://" in url else url.replace("/", "\\")
        r_id = paragraph.part.relate_to(target, RT.HYPERLINK, is_external=True)
        link = OxmlElement("w:hyperlink")
        link.set(qn("r:id"), r_id)
        run = OxmlElement("w:r")
        rPr = OxmlElement("w:rPr")
        col = OxmlElement("w:color"); col.set(qn("w:val"), self.theme[color])
        u = OxmlElement("w:u"); u.set(qn("w:val"), "single")
        sz = OxmlElement("w:sz"); sz.set(qn("w:val"), str(int(size * 2)))
        rPr.append(col); rPr.append(u); rPr.append(sz)
        run.append(rPr)
        t = OxmlElement("w:t"); t.set(qn("xml:space"), "preserve"); t.text = text
        run.append(t)
        link.append(run)
        paragraph._p.append(link)
        return link

    # -- dispatch ----------------------------------------------------------
    def _render(self, doc, mode):
        prefix = "_ws_" if mode == "worksheet" else "_ans_"
        self._title_block(doc, self.meta.get("lesson", "Bài tập"))
        for i, block in enumerate(self.spec.get("blocks", []), start=1):
            btype = block.get("type", "")
            handler = getattr(self, prefix + btype, None)
            if handler is None:
                raise ValueError(
                    "Block #%d: type '%s' không hỗ trợ." % (i, btype))
            handler(doc, block, i)
        return doc

    def render_worksheet(self):
        return self._render(Document(), "worksheet")

    def render_answers(self):
        doc = Document()
        p = doc.add_paragraph()
        self._run(p, "ĐÁP ÁN — ", 13, color="muted", bold=True)
        self._run(p, "chỉ dành cho giáo viên", 13, color="muted", italic=True)
        return self._render(doc, "answers")

    # -- noi (matching) ----------------------------------------------------
    def _ws_noi(self, doc, block, idx):
        self._block_header(doc, idx, block.get("title", "Nối"),
                           block.get("instructions"))
        pairs = block.get("pairs", [])
        rights = self._rotate(list(range(len(pairs))), 1)
        letters = [chr(ord("A") + i) for i in range(len(pairs))]
        table = doc.add_table(rows=len(pairs) + 1, cols=3)
        table.style = "Table Grid"
        hdr = table.rows[0].cells
        self._fill_cell(hdr[0], "汉字", 12, color="accent", bold=True, cjk=True)
        self._fill_cell(hdr[1], "Trả lời", 12, color="accent", bold=True)
        self._fill_cell(hdr[2], "Nghĩa", 12, color="accent", bold=True)
        for r, pair in enumerate(pairs, start=1):
            cells = table.rows[r].cells
            self._fill_cell(cells[0], pair.get("left", ""), 16, bold=True,
                            cjk=True)
            self._fill_cell(cells[1], BLANK, 12, color="muted")
            j = rights[r - 1]
            self._fill_cell(cells[2],
                            "%s. %s" % (letters[r - 1], pairs[j].get("right", "")),
                            12, cjk=True)

    def _ans_noi(self, doc, block, idx):
        self._block_header(doc, idx, block.get("title", "Nối"))
        pairs = block.get("pairs", [])
        rights = self._rotate(list(range(len(pairs))), 1)
        letters = [chr(ord("A") + i) for i in range(len(pairs))]
        pos = {j: letters[k] for k, j in enumerate(rights)}
        for i, pair in enumerate(pairs):
            p = doc.add_paragraph()
            self._run(p, pair.get("left", ""), 14, bold=True, cjk=True)
            self._run(p, "  → %s. " % pos[i], 12, color="accent", bold=True)
            self._run(p, pair.get("right", ""), 12, cjk=True)

    # -- dien_cho_trong (fill blank) --------------------------------------
    def _word_bank(self, doc, words):
        p = doc.add_paragraph()
        self._run(p, "[ ", 12, color="muted", bold=True)
        for i, w in enumerate(words):
            if i:
                self._run(p, " / ", 12, color="muted")
            self._run(p, w, 13, color="accent", bold=True, cjk=True)
        self._run(p, " ]", 12, color="muted", bold=True)

    def _ws_dien_cho_trong(self, doc, block, idx):
        self._block_header(doc, idx, block.get("title", "Điền vào chỗ trống"),
                           block.get("instructions"))
        self._word_bank(doc, block.get("word_bank", []))
        for n, it in enumerate(block.get("items", []), start=1):
            p = doc.add_paragraph()
            self._run(p, "%d) " % n, 12, color="accent", bold=True)
            self._run(p, it.get("q", "").replace("{}", BLANK), 14, cjk=True)

    def _ans_dien_cho_trong(self, doc, block, idx):
        self._block_header(doc, idx, block.get("title", "Điền vào chỗ trống"))
        for n, it in enumerate(block.get("items", []), start=1):
            p = doc.add_paragraph()
            self._run(p, "%d) " % n, 12, color="accent", bold=True)
            self._run(p, it.get("q", "").replace("{}", it.get("answer", "")),
                      14, cjk=True)
            if it.get("src"):
                self._run(p, "  [%s]" % it["src"], 10, color="muted",
                          italic=True)

    # -- doc_hieu (reading MCQ) -------------------------------------------
    def _mcq_questions(self, doc, questions, mark_answer):
        for n, q in enumerate(questions, start=1):
            p = doc.add_paragraph()
            self._run(p, "%d) " % n, 12, color="accent", bold=True)
            self._run(p, q.get("q", ""), 13, cjk=True)
            if not mark_answer:
                self._run(p, "   " + BLANK, 12, color="muted")
            for li, opt in enumerate(q.get("options", [])):
                op = doc.add_paragraph()
                op.paragraph_format.left_indent = Pt(18)
                letter = chr(ord("A") + li)
                is_ans = mark_answer and q.get("answer") == letter
                self._run(op, ("✔ " if is_ans else "") + "%s. " % letter,
                          12, color=("accent" if is_ans else "ink"),
                          bold=is_ans)
                self._run(op, opt, 12, cjk=True, bold=is_ans)
            if mark_answer and q.get("src"):
                sp = doc.add_paragraph()
                self._run(sp, "[%s]" % q["src"], 10, color="muted", italic=True)

    def _ws_doc_hieu(self, doc, block, idx):
        self._block_header(doc, idx, block.get("title", "Đọc hiểu"),
                           block.get("instructions"))
        pp = doc.add_paragraph()
        self._run(pp, block.get("passage", ""), 13, cjk=True)
        self._mcq_questions(doc, block.get("questions", []), mark_answer=False)

    def _ans_doc_hieu(self, doc, block, idx):
        self._block_header(doc, idx, block.get("title", "Đọc hiểu"))
        self._mcq_questions(doc, block.get("questions", []), mark_answer=True)

    # -- sap_xep (sentence ordering) -----------------------------------------
    def _ws_sap_xep(self, doc, block, idx):
        self._block_header(doc, idx, block.get("title", "Sắp xếp câu"),
                           block.get("instructions"))
        for n, it in enumerate(block.get("items", []), start=1):
            chips = self._rotate(list(it.get("words", [])), 2)
            p = doc.add_paragraph()
            self._run(p, "%d) " % n, 12, color="accent", bold=True)
            self._run(p, " / ".join(chips), 14, cjk=True)
            ans = doc.add_paragraph()
            self._run(ans, "→ " + BLANK * 3, 12, color="muted")

    def _two_level_answer(self, doc, n, standard, advanced=None, src=None,
                          alts=None, std_label="Chuẩn (đủ điểm)",
                          adv_label="Nâng cao (điểm cao)"):
        """Đáp án tự luận 2 cấp + (tùy chọn) các phương án khác (alts)."""
        p = doc.add_paragraph()
        self._run(p, "%d) " % n, 12, color="accent", bold=True)
        self._run(p, std_label + ": ", 11, color="muted", bold=True)
        self._run(p, standard, 14, cjk=True)
        if src:
            self._run(p, "  [%s]" % src, 10, color="muted", italic=True)
        if advanced:
            ap = doc.add_paragraph()
            ap.paragraph_format.left_indent = Pt(18)
            self._run(ap, adv_label + ": ", 11, color="accent", bold=True)
            self._run(ap, advanced, 13, cjk=True)
        for alt in (alts or []):
            altp = doc.add_paragraph()
            altp.paragraph_format.left_indent = Pt(18)
            self._run(altp, "Phương án khác: ", 11, color="muted", bold=True)
            self._run(altp, alt, 13, cjk=True)

    def _ans_sap_xep(self, doc, block, idx):
        self._block_header(doc, idx, block.get("title", "Sắp xếp câu"))
        for n, it in enumerate(block.get("items", []), start=1):
            self._two_level_answer(doc, n, it.get("answer", ""),
                                   it.get("answer_plus"), it.get("src"),
                                   alts=it.get("answer_alts"))

    # -- dich_dat_cau (translate / compose) --------------------------------
    def _ws_dich_dat_cau(self, doc, block, idx):
        self._block_header(doc, idx, block.get("title", "Dịch / Đặt câu"),
                           block.get("instructions"))
        for n, it in enumerate(block.get("items", []), start=1):
            p = doc.add_paragraph()
            self._run(p, "%d) " % n, 12, color="accent", bold=True)
            self._run(p, it.get("prompt", ""), 13)
            if it.get("given"):
                self._run(p, "  (dùng: %s)" % " / ".join(it["given"]), 12,
                          color="muted", cjk=True)
            ans = doc.add_paragraph()
            self._run(ans, "→ " + BLANK * 3, 12, color="muted")

    def _ans_dich_dat_cau(self, doc, block, idx):
        self._block_header(doc, idx, block.get("title", "Dịch / Đặt câu"))
        for n, it in enumerate(block.get("items", []), start=1):
            self._two_level_answer(doc, n, it.get("answer", ""),
                                   it.get("answer_plus"), it.get("src"),
                                   alts=it.get("answer_alts"))

    # -- noi_hskk (speaking) ----------------------------------------------
    def _ws_noi_hskk(self, doc, block, idx):
        part = block.get("part", "")
        title = block.get("title", "Nói")
        self._block_header(doc, idx, "🗣 %s (%s)" % (title, part),
                           block.get("instructions"))
        repeat = part == "听后重复"
        for n, it in enumerate(block.get("items", []), start=1):
            p = doc.add_paragraph()
            self._run(p, "%d) " % n, 12, color="accent", bold=True)
            link = it.get("audio_url") or it.get("audio")
            if link:
                self._run(p, "🔊 ", 11, color="accent")
                self._audio_link(p, link, link)
                self._run(p, "   ", 11)
            if repeat:
                # target sentence is the listening stimulus -> hide it
                self._run(p, "(nghe audio rồi nhắc lại)", 12, color="muted",
                          italic=True)
            else:
                self._run(p, it.get("script", ""), 13, cjk=True)

    def _ans_noi_hskk(self, doc, block, idx):
        part = block.get("part", "")
        self._block_header(doc, idx,
                           "🗣 %s (%s)" % (block.get("title", "Nói"), part))
        for n, it in enumerate(block.get("items", []), start=1):
            p = doc.add_paragraph()
            self._run(p, "%d) " % n, 12, color="accent", bold=True)
            self._run(p, it.get("script", ""), 13, cjk=True)
            if it.get("hint"):
                hp = doc.add_paragraph()
                hp.paragraph_format.left_indent = Pt(18)
                self._run(hp, "Gợi ý — Chuẩn (đạt): ", 11, color="muted",
                          bold=True)
                self._run(hp, it["hint"], 12, cjk=True)
            if it.get("hint_plus"):
                hp2 = doc.add_paragraph()
                hp2.paragraph_format.left_indent = Pt(18)
                self._run(hp2, "Gợi ý — Nâng cao (điểm cao): ", 11,
                          color="accent", bold=True)
                self._run(hp2, it["hint_plus"], 12, cjk=True)

    # -- nghe (listening) -------------------------------------------------
    def _ws_nghe(self, doc, block, idx):
        self._block_header(doc, idx, "🔊 " + block.get("title", "Nghe"),
                           block.get("instructions"))
        for n, it in enumerate(block.get("items", []), start=1):
            link = it.get("audio_url") or it.get("audio")
            if link:
                lp = doc.add_paragraph()
                self._run(lp, "🔊 Nghe câu %d: " % n, 11, color="accent",
                          italic=True)
                self._audio_link(lp, link, link)
            qp = doc.add_paragraph()
            self._run(qp, "%d) " % n, 12, color="accent", bold=True)
            self._run(qp, it.get("q", ""), 13, cjk=True)
            self._run(qp, "   " + BLANK, 12, color="muted")
            for li, opt in enumerate(it.get("options", [])):
                op = doc.add_paragraph()
                op.paragraph_format.left_indent = Pt(18)
                self._run(op, "%s. " % chr(ord("A") + li), 12)
                self._run(op, opt, 12, cjk=True)

    def _ans_nghe(self, doc, block, idx):
        self._block_header(doc, idx, "🔊 " + block.get("title", "Nghe"))
        for n, it in enumerate(block.get("items", []), start=1):
            sp = doc.add_paragraph()
            self._run(sp, "%d) 听力文本: " % n, 12, color="muted", bold=True)
            self._run(sp, it.get("script", ""), 13, cjk=True)
            ap = doc.add_paragraph()
            self._run(ap, "✔ %s" % it.get("answer", ""), 12, color="accent",
                      bold=True)
            if it.get("src"):
                self._run(ap, "  [%s]" % it["src"], 10, color="muted",
                          italic=True)

    # -- grammar_note (giải thích ngữ pháp) -------------------------------
    def _grammar_note(self, doc, block, idx):
        self._block_header(doc, idx, block.get("title", "Ngữ pháp"),
                           block.get("instructions"))
        for pt in block.get("points", []):
            p = doc.add_paragraph()
            self._run(p, "• ", 12, color="accent", bold=True)
            self._run(p, pt.get("pattern", ""), 13, color="ink", bold=True,
                      cjk=True)
            if pt.get("explain"):
                ep = doc.add_paragraph()
                ep.paragraph_format.left_indent = Pt(18)
                self._run(ep, pt["explain"], 12, color="ink")
            if pt.get("example"):
                xp = doc.add_paragraph()
                xp.paragraph_format.left_indent = Pt(18)
                self._run(xp, "Vd: ", 11, color="muted", bold=True)
                self._run(xp, pt["example"], 12, cjk=True)

    _ws_grammar_note = _grammar_note
    _ans_grammar_note = _grammar_note

    # -- writing_prompt (đề viết/HSKK + dàn ý) ----------------------------
    # kind: tự do (vd "viết"/"HSKK"/"đoạn văn"/"lời nhắn"/"nhật ký") — không
    # giới hạn danh sách cứng, chỉ hiển thị làm nhãn.
    def _writing_prompt(self, doc, block, idx):
        self._block_header(doc, idx, block.get("title", "Bài viết / HSKK"),
                           block.get("instructions"))
        for n, it in enumerate(block.get("items", []), start=1):
            p = doc.add_paragraph()
            self._run(p, "%d) " % n, 12, color="accent", bold=True)
            if it.get("kind"):
                self._run(p, "[%s] " % it["kind"], 11, color="muted", bold=True)
            self._run(p, it.get("prompt", ""), 13, cjk=True)
            if it.get("target_length"):
                self._run(p, "  (%s)" % it["target_length"], 11, color="muted",
                          italic=True)
            for step in it.get("outline", []):
                sp = doc.add_paragraph()
                sp.paragraph_format.left_indent = Pt(18)
                self._run(sp, "– ", 12, color="muted")
                self._run(sp, step, 12, cjk=True)

    _ws_writing_prompt = _writing_prompt
    _ans_writing_prompt = _writing_prompt

    # -- dien_bieu_mau (điền biểu mẫu — Viết 3.0) --------------------------
    # Form điền thông tin (không có 1 đáp án đúng duy nhất): worksheet in
    # nhãn + dòng trống; đáp án in nhãn + "sample" (câu mẫu tham khảo) nếu có.
    def _ws_dien_bieu_mau(self, doc, block, idx):
        self._block_header(doc, idx, block.get("title", "Điền biểu mẫu"),
                           block.get("instructions"))
        for field in block.get("fields", []):
            p = doc.add_paragraph()
            self._run(p, field.get("label", ""), 13, bold=True, cjk=True)
            self._run(p, "：" + BLANK, 12, color="muted")

    def _ans_dien_bieu_mau(self, doc, block, idx):
        self._block_header(doc, idx, block.get("title", "Điền biểu mẫu"))
        for field in block.get("fields", []):
            p = doc.add_paragraph()
            self._run(p, field.get("label", ""), 13, bold=True, cjk=True)
            if field.get("sample"):
                self._run(p, "：", 12, color="muted")
                self._run(p, field["sample"], 12, color="accent", cjk=True)

    # -- render_study: 1 doc gộp có đáp án (cho lesson-prep) ---------------
    def render_study(self, header_text="BÀI TẬP CHUẨN BỊ"):
        doc = Document()
        p = doc.add_paragraph()
        self._run(p, header_text, 13, color="muted", bold=True)
        return self._render(doc, "answers")


def docx_to_pdf(docx_path):
    """Convert .docx -> .pdf via LibreOffice if available; else return None."""
    soffice = shutil.which("soffice") or shutil.which("soffice.exe")
    if not soffice:
        return None
    src = Path(docx_path)
    out_dir = src.parent
    try:
        subprocess.run(
            [soffice, "--headless", "--convert-to", "pdf", "--outdir",
             str(out_dir), str(src)],
            check=True, capture_output=True, timeout=120)
    except (subprocess.SubprocessError, OSError):
        return None
    pdf = src.with_suffix(".pdf")
    return str(pdf) if pdf.exists() else None


def main(argv):
    # Console trên Windows mặc định là cp1252; các dòng thông báo có tiếng Việt
    # sẽ crash khi in. Ép stdout/stderr sang UTF-8 (bỏ qua nếu không hỗ trợ).
    for stream in (sys.stdout, sys.stderr):
        try:
            stream.reconfigure(encoding="utf-8")
        except (AttributeError, ValueError):
            pass
    if len(argv) != 3:
        print("Usage: python build_worksheet.py <baitap.json> <out_dir>",
              file=sys.stderr)
        return 2
    src = Path(argv[1])
    out_dir = Path(argv[2])
    # Tách output: phần đưa học sinh (worksheet + audio) vào hocsinh/, đáp án
    # vào dapan/ → gửi cả folder hocsinh/ cho học sinh mà không lộ đáp án.
    # Link audio trong worksheet là tương đối (audio/...), nên folder audio/
    # PHẢI nằm trong hocsinh/ cạnh worksheet.docx.
    hocsinh_dir = out_dir / "hocsinh"
    dapan_dir = out_dir / "dapan"
    hocsinh_dir.mkdir(parents=True, exist_ok=True)
    dapan_dir.mkdir(parents=True, exist_ok=True)
    spec = json.loads(src.read_text(encoding="utf-8"))
    b = WorksheetBuilder(spec)
    b.render_worksheet().save(str(hocsinh_dir / "worksheet.docx"))
    b.render_answers().save(str(dapan_dir / "dapan.docx"))
    made = []
    for path in (hocsinh_dir / "worksheet.docx", dapan_dir / "dapan.docx"):
        pdf = docx_to_pdf(str(path))
        if pdf:
            made.append(Path(pdf).name)
    n = len(spec.get("blocks", []))
    print("OK: %d block -> hocsinh/worksheet.docx + dapan/dapan.docx (%s)"
          % (n, out_dir))
    if made:
        print("PDF: " + ", ".join(made))
    else:
        print("PDF: chưa xuất (không thấy LibreOffice). "
              "Mở .docx và 'Save as PDF' nếu cần.")
    return 0


if __name__ == "__main__":
    raise SystemExit(main(sys.argv))
