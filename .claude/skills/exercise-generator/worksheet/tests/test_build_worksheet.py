from docx.oxml.ns import qn
from build_worksheet import WorksheetBuilder


def _run_has_cjk_font(run, font="Microsoft YaHei"):
    rPr = run._element.rPr
    if rPr is None or rPr.rFonts is None:
        return False
    return rPr.rFonts.get(qn("w:eastAsia")) == font


def _all_text(doc):
    parts = [p.text for p in doc.paragraphs]
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "\n".join(parts)


def test_core_renders_lesson_title_with_cjk():
    spec = {"meta": {"lesson": "Buổi 1: 会/想/能", "hsk": 1}, "blocks": []}
    doc = WorksheetBuilder(spec).render_worksheet()
    assert "Buổi 1" in _all_text(doc)
    # The title run carrying Chinese must declare the East-Asian font.
    title_runs = [r for p in doc.paragraphs for r in p.runs if "会" in r.text]
    assert title_runs and _run_has_cjk_font(title_runs[0])


def test_core_unknown_block_type_raises():
    spec = {"meta": {"lesson": "x", "hsk": 1}, "blocks": [{"type": "nope"}]}
    try:
        WorksheetBuilder(spec).render_worksheet()
    except ValueError as e:
        assert "nope" in str(e)
    else:
        raise AssertionError("expected ValueError for unknown block type")


NOI_BLOCK = {
    "type": "noi",
    "title": "Nối chữ với nghĩa",
    "instructions": "Viết chữ cái đúng vào ô.",
    "pairs": [
        {"left": "会", "right": "biết (kỹ năng)"},
        {"left": "想", "right": "muốn"},
        {"left": "能", "right": "có thể"},
    ],
}


def _spec(*blocks):
    return {"meta": {"lesson": "Buổi 1", "hsk": 1}, "blocks": list(blocks)}


def test_noi_worksheet_lists_all_left_items_and_no_direct_answer_pairing():
    doc = WorksheetBuilder(_spec(NOI_BLOCK)).render_worksheet()
    text = _all_text(doc)
    for hz in ("会", "想", "能"):
        assert hz in text
    # meanings are present (as a shuffled lettered bank), instruction shown
    assert "muốn" in text
    assert "Viết chữ cái đúng" in text


def test_noi_shuffle_never_pairs_character_with_own_meaning():
    doc = WorksheetBuilder(_spec(NOI_BLOCK)).render_worksheet()
    table = doc.tables[0]
    correct_meaning = {p["left"]: p["right"] for p in NOI_BLOCK["pairs"]}
    for row in table.rows[1:]:  # skip header row
        cells = row.cells
        hz = cells[0].text.strip()
        shown = cells[2].text.split(". ", 1)[-1].strip()
        assert shown != correct_meaning[hz], (
            "row for %r shows its own correct meaning %r" % (hz, shown))


def test_noi_answer_key_shows_mapping():
    doc = WorksheetBuilder(_spec(NOI_BLOCK)).render_answers()
    text = _all_text(doc)
    # answer key pairs each character with its meaning
    assert "会" in text and "biết (kỹ năng)" in text


FILL_BLOCK = {
    "type": "dien_cho_trong",
    "title": "Điền từ",
    "instructions": "Chọn từ trong khung.",
    "word_bank": ["会", "想", "能"],
    "items": [
        {"q": "我{}说一点儿汉语。", "answer": "会"},
        {"q": "今天我不{}去。", "answer": "能", "src": "phỏng theo 真题"},
    ],
}


def test_fill_worksheet_shows_blanks_and_bank_but_no_answers():
    from build_worksheet import BLANK
    doc = WorksheetBuilder(_spec(FILL_BLOCK)).render_worksheet()
    text = _all_text(doc)
    assert "会" in text and "想" in text  # bank present
    assert BLANK in text                   # blanks rendered
    # The full answered sentence "我会说一点儿汉语。" must NOT appear.
    assert "我会说一点儿汉语。" not in text


def test_fill_answer_key_fills_and_labels_generated():
    doc = WorksheetBuilder(_spec(FILL_BLOCK)).render_answers()
    text = _all_text(doc)
    assert "我会说一点儿汉语。" in text
    assert "phỏng theo 真题" in text


READ_BLOCK = {
    "type": "doc_hieu",
    "title": "Đọc hiểu",
    "passage": "小明会说汉语，也会说英语。",
    "questions": [
        {"q": "小明会说什么？", "options": ["汉语和英语", "只有汉语", "日语"],
         "answer": "A"},
    ],
}


def test_read_worksheet_shows_passage_options_no_marked_answer():
    doc = WorksheetBuilder(_spec(READ_BLOCK)).render_worksheet()
    text = _all_text(doc)
    assert "小明会说汉语" in text
    assert "A." in text and "汉语和英语" in text
    assert "✔" not in text  # no answer marker on the worksheet


def test_read_answer_key_marks_correct_letter():
    doc = WorksheetBuilder(_spec(READ_BLOCK)).render_answers()
    text = _all_text(doc)
    assert "✔ A" in text


ORDER_BLOCK = {
    "type": "sap_xep",
    "title": "Sắp xếp câu",
    "items": [{"words": ["我", "会", "说", "汉语"], "answer": "我会说汉语。"}],
}


def test_order_worksheet_shows_chips_not_answer():
    doc = WorksheetBuilder(_spec(ORDER_BLOCK)).render_worksheet()
    text = _all_text(doc)
    assert "会" in text and "汉语" in text
    assert "我会说汉语。" not in text  # answer hidden


def test_order_answer_key_shows_sentence():
    doc = WorksheetBuilder(_spec(ORDER_BLOCK)).render_answers()
    assert "我会说汉语。" in _all_text(doc)


TRANS_BLOCK = {
    "type": "dich_dat_cau",
    "title": "Dịch câu",
    "items": [{"prompt": "Tôi muốn học tiếng Trung.", "given": ["想", "学"],
               "answer": "我想学汉语。"}],
}


def test_trans_worksheet_shows_prompt_and_given_not_answer():
    doc = WorksheetBuilder(_spec(TRANS_BLOCK)).render_worksheet()
    text = _all_text(doc)
    assert "Tôi muốn học tiếng Trung." in text
    assert "想" in text and "学" in text
    assert "我想学汉语。" not in text


def test_trans_answer_key_shows_model_sentence():
    doc = WorksheetBuilder(_spec(TRANS_BLOCK)).render_answers()
    assert "我想学汉语。" in _all_text(doc)


LISTEN_BLOCK = {
    "type": "nghe",
    "title": "Nghe và chọn",
    "instructions": "Nghe rồi chọn đáp án đúng.",
    "items": [
        {"script": "你会游泳吗？", "q": "对话问什么？",
         "options": ["会不会游泳", "会不会开车", "想不想吃饭"], "answer": "A",
         "audio": "audio/nghe-1.mp3"},
    ],
}


def test_listen_worksheet_hides_script_shows_question():
    doc = WorksheetBuilder(_spec(LISTEN_BLOCK)).render_worksheet()
    text = _all_text(doc)
    assert "你会游泳吗？" not in text          # script hidden on worksheet
    assert "对话问什么？" in text               # question shown
    assert "会不会游泳" in text                 # options shown
    assert "audio/nghe-1.mp3" in text          # audio link shown


def test_listen_answer_key_reveals_script_and_answer():
    doc = WorksheetBuilder(_spec(LISTEN_BLOCK)).render_answers()
    text = _all_text(doc)
    assert "你会游泳吗？" in text               # 听力文本 revealed
    assert "✔ A" in text


SPEAK_REPEAT = {
    "type": "noi_hskk", "title": "Nghe và nhắc lại", "part": "听后重复",
    "items": [{"script": "我会说一点儿汉语。", "audio": "audio/hskk-1.mp3"}],
}
SPEAK_ANSWER = {
    "type": "noi_hskk", "title": "Trả lời câu hỏi", "part": "回答问题",
    "items": [{"script": "你会做饭吗？", "hint": "我会做饭。/ 我不会做饭。"}],
}


def test_speak_repeat_hides_target_on_worksheet():
    doc = WorksheetBuilder(_spec(SPEAK_REPEAT)).render_worksheet()
    text = _all_text(doc)
    assert "听后重复" in text
    assert "我会说一点儿汉语。" not in text  # target hidden for repeat drill
    assert "audio/hskk-1.mp3" in text


def test_speak_answer_shows_question_on_worksheet_hint_in_key():
    ws = _all_text(WorksheetBuilder(_spec(SPEAK_ANSWER)).render_worksheet())
    ans = _all_text(WorksheetBuilder(_spec(SPEAK_ANSWER)).render_answers())
    assert "你会做饭吗？" in ws           # question is shown for 回答问题
    assert "我会做饭。" not in ws          # suggested answer hidden on worksheet
    assert "我会做饭。" in ans             # suggested answer in the key


import json
import pathlib
from build_worksheet import main
from audio_manifest import build_audio_manifest
import shutil
from build_worksheet import docx_to_pdf


def test_audio_manifest_collects_listening_and_speaking_jobs():
    spec = _spec(LISTEN_BLOCK, SPEAK_REPEAT, TRANS_BLOCK)
    jobs = build_audio_manifest(spec)
    files = {j["file"] for j in jobs}
    assert files == {"audio/nghe-1.mp3", "audio/hskk-1.mp3"}
    nghe = next(j for j in jobs if j["file"] == "audio/nghe-1.mp3")
    assert nghe["text"] == "你会游泳吗？"
    assert nghe["voice"] == "zh-CN-XiaoxiaoNeural"


def test_pdf_fallback_returns_none_when_no_converter(tmp_path, monkeypatch):
    # Simulate no LibreOffice on PATH.
    monkeypatch.setattr(shutil, "which", lambda name: None)
    fake = tmp_path / "worksheet.docx"
    fake.write_bytes(b"not a real docx")
    assert docx_to_pdf(str(fake)) is None


ALL_BLOCKS = [NOI_BLOCK, FILL_BLOCK, READ_BLOCK, ORDER_BLOCK, TRANS_BLOCK,
              LISTEN_BLOCK, SPEAK_REPEAT, SPEAK_ANSWER]


def test_no_listening_script_leaks_into_worksheet():
    spec = {"meta": {"lesson": "Buổi 1: 会/想/能", "hsk": 1},
            "blocks": ALL_BLOCKS}
    ws = _all_text(WorksheetBuilder(spec).render_worksheet())
    # Every listening / repeat script must be absent from the worksheet.
    for hidden in ("你会游泳吗？", "我会说一点儿汉语。"):
        assert hidden not in ws
    # But the answer key must contain them.
    key = _all_text(WorksheetBuilder(spec).render_answers())
    assert "你会游泳吗？" in key and "我会说一点儿汉语。" in key


FIXTURE = pathlib.Path(__file__).resolve().parent.parent / "example-baitap.json"


def test_example_fixture_has_all_seven_block_types():
    spec = json.loads(FIXTURE.read_text(encoding="utf-8"))
    types = {b["type"] for b in spec["blocks"]}
    assert types == {"noi", "dien_cho_trong", "doc_hieu", "sap_xep",
                     "dich_dat_cau", "nghe", "noi_hskk"}


def test_end_to_end_cli_writes_both_docs(tmp_path):
    rc = main(["build_worksheet.py", str(FIXTURE), str(tmp_path)])
    assert rc == 0
    # Output tách: worksheet + audio vào hocsinh/, đáp án vào dapan/
    assert (tmp_path / "hocsinh" / "worksheet.docx").exists()
    assert (tmp_path / "dapan" / "dapan.docx").exists()
    # reopen worksheet and confirm CJK + no listening-script leak
    from docx import Document
    ws = _all_text(Document(str(tmp_path / "hocsinh" / "worksheet.docx")))
    assert "会" in ws
    assert "你会游泳吗？" not in ws


# --- 2-level answer key (standard / advanced) for productive blocks ---------
DICH_2LEVEL = {
    "type": "dich_dat_cau", "title": "Dịch",
    "items": [{"prompt": "Mình biết bơi.", "given": ["会"],
               "answer": "我会游泳。",
               "answer_plus": "我会游泳，我每天都去游泳池。",
               "src": "phỏng theo 真题"}],
}
SAPXEP_2LEVEL = {
    "type": "sap_xep", "title": "Sắp xếp",
    "items": [{"words": ["我", "会", "做", "中国菜"], "answer": "我会做中国菜。",
               "answer_plus": "我会做中国菜，也会做越南菜。"}],
}
HSKK_2LEVEL = {
    "type": "noi_hskk", "title": "Trả lời", "part": "回答问题",
    "items": [{"script": "你会做什么菜？",
               "hint": "我会做西红柿炒鸡蛋。",
               "hint_plus": "我会做几个中国菜，比如西红柿炒鸡蛋，周末常常做。"}],
}


def test_dich_answer_key_shows_two_levels():
    ans = _all_text(WorksheetBuilder(_spec(DICH_2LEVEL)).render_answers())
    assert "我会游泳。" in ans
    assert "我会游泳，我每天都去游泳池。" in ans
    assert "Chuẩn" in ans and "Nâng cao" in ans


def test_dich_worksheet_hides_both_levels():
    ws = _all_text(WorksheetBuilder(_spec(DICH_2LEVEL)).render_worksheet())
    assert "我会游泳。" not in ws
    assert "我会游泳，我每天都去游泳池。" not in ws


def test_dich_answer_without_plus_shows_single_no_advanced_line():
    single = {"type": "dich_dat_cau", "title": "Dịch",
              "items": [{"prompt": "x", "answer": "我会游泳。"}]}
    ans = _all_text(WorksheetBuilder(_spec(single)).render_answers())
    assert "我会游泳。" in ans
    assert "Nâng cao" not in ans


def test_sapxep_answer_key_shows_two_levels():
    ans = _all_text(WorksheetBuilder(_spec(SAPXEP_2LEVEL)).render_answers())
    assert "我会做中国菜。" in ans
    assert "我会做中国菜，也会做越南菜。" in ans
    assert "Nâng cao" in ans


def test_hskk_answer_key_shows_two_level_hints():
    ans = _all_text(WorksheetBuilder(_spec(HSKK_2LEVEL)).render_answers())
    assert "我会做西红柿炒鸡蛋。" in ans
    assert "我会做几个中国菜" in ans
    assert "Chuẩn" in ans and "Nâng cao" in ans


# --- audio links must be clickable hyperlinks in the worksheet --------------
def _external_hyperlink_targets(doc):
    part = doc.part
    return [rel.target_ref for rel in part.rels.values()
            if rel.reltype.endswith("hyperlink") and rel.is_external]


def test_nghe_worksheet_audio_is_clickable_hyperlink():
    # local file targets use Windows backslashes so Word opens them as files
    doc = WorksheetBuilder(_spec(LISTEN_BLOCK)).render_worksheet()
    assert "audio\\nghe-1.mp3" in _external_hyperlink_targets(doc)


def test_hskk_repeat_audio_is_clickable_hyperlink():
    doc = WorksheetBuilder(_spec(SPEAK_REPEAT)).render_worksheet()
    assert "audio\\hskk-1.mp3" in _external_hyperlink_targets(doc)


def test_audio_url_when_present_wins_over_local_path():
    block = {"type": "nghe", "title": "Nghe", "items": [
        {"script": "你好？", "q": "?", "options": ["a", "b"], "answer": "A",
         "audio": "audio/x.mp3", "audio_url": "https://ex.com/x.mp3"}]}
    doc = WorksheetBuilder(_spec(block)).render_worksheet()
    assert "https://ex.com/x.mp3" in _external_hyperlink_targets(doc)
