# -*- coding: utf-8 -*-
"""render_baitap.py — render bài tập lesson-prep theo MẪU của user (raw/bai tap mau.docx).

Thứ tự & quy tắc:
  1) Bài viết lớn (作文) xếp ĐẦU — chỉ đề, không giải.
  2) Bỏ qua luyện đọc từ / đặt câu với từ.
  3) Mỗi đề: "Đáp án:" của bài; riêng câu hoàn thành (type=complete) thêm Opt1/2/3
     (câu khác tự đặt); 改写/đục lỗ/判断 (1 đáp án) thì chỉ 1 đáp án.

Schema JSON:
{
  "meta": {"lesson": "..."},
  "writing": {"title": "Bài viết lớn (作文)", "prompt": "..."},
  "exercises": [
    {"title": "...", "type": "complete|rewrite|judge|cloze", "note": "?",
     "items": [{"q": "...", "answer": "...", "opts": ["..."], "vi": "..."}]}
  ]
}
Chạy:  python render_baitap.py <exercise_payload.json> <out.docx>
Cần: python-docx.
"""
import sys, json, re
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn

CJK = "Microsoft YaHei"
TEXT = "Calibri"
ACCENT = "C0392B"
MUTED = "6B7683"
INK = "1F2933"
RED = "EE0000"  # tô từ/cấu trúc trọng tâm trong đáp án (như file mẫu)


def run(p, text, size=12, color=INK, bold=False, italic=False):
    r = p.add_run(text)
    f = r.font
    f.size = Pt(size); f.bold = bold; f.italic = italic
    f.color.rgb = RGBColor.from_string(color)
    f.name = TEXT
    r._element.rPr.rFonts.set(qn("w:eastAsia"), CJK)
    return r


def run_hl(p, text, hls, size=12.5):
    """Tô ĐỎ các chuỗi trong hls (từ/cấu trúc trọng tâm), phần còn lại màu thường."""
    hls = [h for h in (hls or []) if h]
    if not hls:
        run(p, text, size); return
    pat = "(" + "|".join(re.escape(h) for h in hls) + ")"
    for seg in re.split(pat, text):
        if not seg:
            continue
        if seg in hls:
            run(p, seg, size, color=RED, bold=True)
        else:
            run(p, seg, size)


def para(doc, before=0, after=0, indent=0):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    if indent:
        p.paragraph_format.left_indent = Pt(indent)
    return p


def render(spec, out_path):
    doc = Document()
    meta = spec.get("meta", {})
    if meta.get("lesson"):
        p = para(doc, after=8)
        run(p, meta["lesson"], 15, color=ACCENT, bold=True)

    # 1) Bài viết lớn — chỉ đề
    w = spec.get("writing")
    if w:
        p = para(doc, before=2, after=2)
        run(p, w.get("title", "Bài viết lớn (作文)"), 14, color=ACCENT, bold=True)
        pp = para(doc, after=6)
        run(pp, "Đề: ", 12, color=MUTED, bold=True)
        run(pp, w.get("prompt", ""), 12.5)

    # 2) Các đề bài tập
    for i, ex in enumerate(spec.get("exercises", []), start=1):
        p = para(doc, before=12, after=2)
        run(p, "%d. " % i, 13, color=ACCENT, bold=True)
        run(p, ex.get("title", ""), 13, bold=True)
        if ex.get("note"):
            np = para(doc, after=2)
            run(np, ex["note"], 10.5, color=MUTED, italic=True)
        etype = ex.get("type", "")
        for it in ex.get("items", []):
            q = para(doc, before=4)
            run(q, it.get("q", ""), 12.5)
            if it.get("vi"):
                run(q, "  (%s)" % it["vi"], 11, color=MUTED, italic=True)
            hl = it.get("hl", [])
            if etype == "rewrite" or etype == "judge":
                ap = para(doc, indent=18)
                run(ap, "→ ", 12, color=ACCENT, bold=True)
                run_hl(ap, it.get("answer", ""), hl)
            else:  # complete, cloze
                ap = para(doc, indent=18)
                run(ap, "Đáp án: ", 11.5, color=ACCENT, bold=True)
                run_hl(ap, it.get("answer", ""), hl)
                for k, opt in enumerate(it.get("opts", []), start=1):
                    op = para(doc, indent=18)
                    run(op, "Opt%d: " % k, 11, color=MUTED, bold=True)
                    run_hl(op, opt, hl)
    Path(out_path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    return len(spec.get("exercises", []))


def main(argv):
    for s in (sys.stdout, sys.stderr):
        try:
            s.reconfigure(encoding="utf-8")
        except (AttributeError, ValueError):
            pass
    if len(argv) != 3:
        print("Usage: python render_baitap.py <exercise_payload.json> <out.docx>",
              file=sys.stderr); return 2
    spec = json.loads(Path(argv[1]).read_text(encoding="utf-8"))
    n = render(spec, argv[2])
    print("OK: baitap -> %s (%d đề + bài viết)" % (argv[2], n))
    return 0


if __name__ == "__main__":
    raise SystemExit(main(sys.argv))
